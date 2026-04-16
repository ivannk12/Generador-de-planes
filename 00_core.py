# ==============================================
# 📚 GENERADOR DE PLANES ICFES — PREICFES MATERIAL
# Core reutilizable + 3 runners (Tipo 1/2/3)
#
# ✅ Diseñado para Google Colab + python-docx
# ✅ Mantiene diseño visual (header con logos, portada, tabla semanal 2×7, colores, checkbox)
# ✅ Un solo motor base + estrategias por tipo (sin duplicar render)
#
# 👉 USO EN COLAB (recomendado):
#   1) Copia este archivo en una celda llamada "CORE" y ejecútala.
#   2) Copia SOLO el runner del tipo que necesites en otra celda y ejecútalo.
#
# NOTA: En Colab, sube el logo con files.upload() y pon el nombre del archivo en branding.logo_path.
# ==============================================

# =========================
# ======= CELL: CORE ======
# =========================

from __future__ import annotations

from dataclasses import dataclass
from datetime import date, datetime, timedelta
from typing import Dict, List, Optional, Tuple, Any
import math

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Colab helpers (si no estás en Colab, no pasa nada)
try:
    from google.colab import files
except Exception:
    files = None


# ------------------------------------------------
# 0) Constantes ICFES (estructura fija del cuadernillo)
# ------------------------------------------------

# Orden lineal REAL dentro del cuadernillo (para Tipo 1)
BOOKLET_SLOTS: List[Tuple[str, int]] = [
    ("Matemáticas", 25),
    ("Lectura Crítica", 41),
    ("Sociales y Ciudadanas", 25),
    ("Ciencias Naturales", 29),
    ("Sociales y Ciudadanas", 25),
    ("Matemáticas", 25),
    ("Ciencias Naturales", 29),
    ("Inglés", 55),
]

# Totales por materia (derivados de los slots)
SUBJECT_TOTALS: Dict[str, int] = {}
for s, n in BOOKLET_SLOTS:
    SUBJECT_TOTALS[s] = SUBJECT_TOTALS.get(s, 0) + n

TOTAL_BOOKLET_Q = sum(n for _, n in BOOKLET_SLOTS)  # 254

DAYS_ES = ["lunes", "martes", "miércoles", "jueves", "viernes", "sábado", "domingo"]
DAYNAME_TO_WEEKDAY = {
    "lunes": 0,
    "martes": 1,
    "miercoles": 2,
    "miércoles": 2,
    "jueves": 3,
    "viernes": 4,
    "sabado": 5,
    "sábado": 5,
    "domingo": 6,
}


# ------------------------------------------------
# 1) Utilidades pequeñas (limpias y testeables)
# ------------------------------------------------

def _round_half_up(x: float) -> int:
    """Redondeo consistente: 22.5 -> 23, 22.49 -> 22"""
    return int(math.floor(x + 0.5))


def hex_to_rgb(hex_color: str) -> RGBColor:
    hex_color = hex_color.strip().lstrip("#")
    if len(hex_color) != 6:
        raise ValueError(f"Color inválido: {hex_color}")
    r = int(hex_color[0:2], 16)
    g = int(hex_color[2:4], 16)
    b = int(hex_color[4:6], 16)
    return RGBColor(r, g, b)


def shade_cell(cell, fill_hex_no_hash: str) -> None:
    """Pinta el fondo de una celda con HEX SIN '#', p.ej. 'A8D8FF'."""
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex_no_hash)
    cell._tc.get_or_add_tcPr().append(shd)


def iter_weeks_monday_to_sunday(start: date, end: date):
    """Genera semanas como listas de 7 fechas (lunes..domingo) cubriendo [start,end]."""
    cursor = start
    while cursor <= end:
        monday = cursor - timedelta(days=cursor.weekday())
        week = [monday + timedelta(days=i) for i in range(7)]
        yield week
        cursor = monday + timedelta(days=7)


def cycle_map(keys: List[str], palette: List[str]) -> Dict[str, str]:
    """Asigna colores cíclicos a una lista de keys."""
    if not palette:
        raise ValueError("palette no puede estar vacía")
    out = {}
    for i, k in enumerate(keys):
        out[k] = palette[i % len(palette)]
    return out


# ------------------------------------------------
# 2) Modelo intermedio (motor -> render)
# ------------------------------------------------

@dataclass
class DayBlock:
    subject: str
    booklet_name: str
    start_q: int
    end_q: int
    qty: int
    range_scope: str = "subject"  # "subject" (Tipo 2/3) | "booklet" (Tipo 1)


@dataclass
class PlanDay:
    day: date
    weekday: int
    is_rest: bool
    color: Optional[str]  # HEX con '#'
    blocks: List[DayBlock]
    notes: List[str]


@dataclass
class PlanModel:
    meta: Dict[str, Any]
    days: List[PlanDay]


# ------------------------------------------------
# 3) Validación mínima del config (sin burocracia)
# ------------------------------------------------

def validate_config(cfg: Dict[str, Any]) -> None:
    # meta
    plan_type = int(cfg["meta"]["plan_type"])
    if plan_type not in (1, 2, 3):
        raise ValueError("meta.plan_type debe ser 1, 2 o 3")

    # fechas
    start = date.fromisoformat(cfg["date_range"]["start"])
    end = date.fromisoformat(cfg["date_range"]["end"])
    if start > end:
        raise ValueError("date_range.start no puede ser mayor que date_range.end")

    # weekdays
    aw = cfg["calendar_rules"]["active_weekdays"]
    if any((d < 0 or d > 6) for d in aw):
        raise ValueError("calendar_rules.active_weekdays solo admite 0..6")

    # palette
    palette = cfg["style"]["palette"]
    if not (1 <= len(palette) <= 5):
        raise ValueError("style.palette debe tener entre 1 y 5 colores")

    # Tipo 1 requiere global_booklets
    if plan_type == 1:
        gb = cfg["content"]["global_booklets"]
        if not gb:
            raise ValueError("content.global_booklets no puede estar vacío para Tipo 1")

    # Tipo 2 requiere per_subject_base_q + multiplier
    if plan_type == 2:
        if not cfg["intensity"].get("per_subject_base_q"):
            raise ValueError("intensity.per_subject_base_q es obligatorio para Tipo 2")
        if not cfg["intensity"].get("multiplier_by_weekday"):
            raise ValueError("intensity.multiplier_by_weekday es obligatorio para Tipo 2")

    # Tipo 3 requiere assignments
    if plan_type == 3:
        if not cfg.get("assignments", {}).get("subject_by_weekday"):
            raise ValueError("assignments.subject_by_weekday es obligatorio para Tipo 3")


# ------------------------------------------------
# 4) Estado de progreso (core reutilizable)
# ------------------------------------------------

@dataclass
class LinearBookletProgress:
    """Progreso de un cuadernillo siguiendo BOOKLET_SLOTS linealmente (Tipo 1)."""
    slot_idx: int = 0
    slot_offset: int = 0
    done_by_subject: Dict[str, int] = None

    def __post_init__(self):
        if self.done_by_subject is None:
            self.done_by_subject = {s: 0 for s in SUBJECT_TOTALS.keys()}

    def is_complete(self) -> bool:
        return self.slot_idx >= len(BOOKLET_SLOTS)


@dataclass
class SubjectProgress:
    """Progreso por materia a través de una secuencia de cuadernillos (Tipo 2 y 3)."""
    booklet_idx: int = 0
    pos_in_booklet: int = 0  # 0..total_q


# ------------------------------------------------
# 5) Cálculo de intensidad diaria
# ------------------------------------------------

def daily_quota_type1_or_3(cfg: Dict[str, Any], weekday: int) -> int:
    mode = cfg["intensity"]["mode"]
    if mode == "fixed":
        return int(cfg["intensity"]["fixed_q"])
    elif mode == "by_weekday":
        m = cfg["intensity"]["by_weekday_q"]
        return int(m.get(str(weekday), 0))
    else:
        raise ValueError("intensity.mode debe ser 'fixed' o 'by_weekday'")


# ------------------------------------------------
# 6) Motores (3 estrategias) -> PlanModel
# ------------------------------------------------


def build_plan_model(cfg: Dict[str, Any]) -> PlanModel:
    validate_config(cfg)
    plan_type = int(cfg["meta"]["plan_type"])

    if plan_type == 1:
        return _build_type1(cfg)
    if plan_type == 2:
        return _build_type2(cfg)
    if plan_type == 3:
        return _build_type3(cfg)

    raise ValueError("plan_type no soportado")


# -------- Tipo 1: cuadernillos en orden, lineal por slots --------

def _build_type1(cfg: Dict[str, Any]) -> PlanModel:
    start = date.fromisoformat(cfg["date_range"]["start"])
    end = date.fromisoformat(cfg["date_range"]["end"])
    active_weekdays = set(cfg["calendar_rules"]["active_weekdays"])

    global_booklets = cfg["content"]["global_booklets"]  # lista dicts: {id,name}
    booklet_names = [b["name"] for b in global_booklets]

    palette = cfg["style"]["palette"]
    booklet_color_map = cycle_map(booklet_names, palette)

    days: List[PlanDay] = []

    booklet_idx = 0
    prog = LinearBookletProgress()

    d = start
    while d <= end and booklet_idx < len(global_booklets):
        wd = d.weekday()
        quota = daily_quota_type1_or_3(cfg, wd)

        # Día fuera de días activos => celda vacía (descanso) sin consumo
        if wd not in active_weekdays:
            days.append(PlanDay(day=d, weekday=wd, is_rest=True, color=None, blocks=[], notes=["Descanso"],))
            d += timedelta(days=1)
            continue

        # Quota 0 => descanso explícito
        if quota <= 0:
            days.append(PlanDay(day=d, weekday=wd, is_rest=True, color=None, blocks=[], notes=["Descanso"],))
            d += timedelta(days=1)
            continue

        blocks: List[DayBlock] = []
        notes: List[str] = []
        remaining = quota

        # Color del día (por cuadernillo actual)
        current_booklet_name = global_booklets[booklet_idx]["name"]
        day_color = booklet_color_map[current_booklet_name]

        while remaining > 0 and booklet_idx < len(global_booklets):
            # si el cuadernillo está completo, avanzar
            if prog.is_complete():
                notes.append(f"🎉 ¡Cuadernillo {current_booklet_name} COMPLETADO! 🎉")
                booklet_idx += 1
                if booklet_idx >= len(global_booklets):
                    break
                # nuevo cuadernillo
                current_booklet_name = global_booklets[booklet_idx]["name"]
                day_color = booklet_color_map[current_booklet_name]
                prog = LinearBookletProgress()
                notes.append(f"📖 Inicio del Cuadernillo {current_booklet_name} 📖")

            if booklet_idx >= len(global_booklets):
                break

            # slot actual
            subject, slot_len = BOOKLET_SLOTS[prog.slot_idx]
            slot_remaining = slot_len - prog.slot_offset
            if slot_remaining <= 0:
                prog.slot_idx += 1
                prog.slot_offset = 0
                continue

            take = min(remaining, slot_remaining)

            # ✅ RANGO CORRECTO PARA TIPO 1 (según tu regla real):
            # La numeración se reinicia por sesión.
            # Sesión 1: 120 preguntas (slots 0..3) => 1..120
            # Sesión 2: 134 preguntas (slots 4..7) => 1..134
            if prog.slot_idx <= 3:
                # sesión 1
                session_before_slot = sum(n for _, n in BOOKLET_SLOTS[:prog.slot_idx])
            else:
                # sesión 2
                session_before_slot = sum(n for _, n in BOOKLET_SLOTS[4:prog.slot_idx])

            start_q = session_before_slot + prog.slot_offset + 1
            end_q = start_q + take - 1

            blocks.append(DayBlock(
                subject=subject,
                booklet_name=current_booklet_name,
                start_q=start_q,
                end_q=end_q,
                qty=take,
                range_scope="booklet",
            ))

            # actualizar progreso lineal
            prog.slot_offset += take
            remaining -= take

            # si se acabó el slot, avanzar
            if prog.slot_offset >= slot_len:
                prog.slot_idx += 1
                prog.slot_offset = 0

        days.append(PlanDay(day=d, weekday=wd, is_rest=False, color=day_color, blocks=blocks, notes=notes))
        d += timedelta(days=1)

    return PlanModel(meta=cfg["meta"], days=days)


# -------- Tipo 2: por materias con multiplicador por día --------

def _build_type2(cfg: Dict[str, Any]) -> PlanModel:
    start = date.fromisoformat(cfg["date_range"]["start"])
    end = date.fromisoformat(cfg["date_range"]["end"])
    active_weekdays = set(cfg["calendar_rules"]["active_weekdays"])

    per_subject_base = cfg["intensity"]["per_subject_base_q"]  # {materia: base}
    mult_map = cfg["intensity"]["multiplier_by_weekday"]        # {"0": 1, ...}

    # Catálogo por materia: cada materia tiene lista de cuadernillos (name/id) y total_q para esa materia
    subj_catalog: Dict[str, List[Dict[str, Any]]] = cfg["content"]["subjects"]

    subjects = [s for s in per_subject_base.keys() if s in subj_catalog]

    palette = cfg["style"]["palette"]
    subject_color_map = cfg["style"].get("subject_color_map") or cycle_map(subjects, palette)

    prog_by_subject: Dict[str, SubjectProgress] = {s: SubjectProgress() for s in subjects}

    days: List[PlanDay] = []
    d = start

    while d <= end:
        wd = d.weekday()

        if wd not in active_weekdays:
            days.append(PlanDay(day=d, weekday=wd, is_rest=True, color=None, blocks=[], notes=["Descanso"],))
            d += timedelta(days=1)
            continue

        mult = float(mult_map.get(str(wd), 1))
        if mult == 0:
            # Confirmado: descanso total, sin bloques
            days.append(PlanDay(day=d, weekday=wd, is_rest=True, color=None, blocks=[], notes=["Descanso"],))
            d += timedelta(days=1)
            continue

        blocks: List[DayBlock] = []
        notes: List[str] = []

        # En Tipo 2 el color "del día" no es único. Para no romper tu diseño (1 color por celda),
        # usamos un color neutro (primer color) o None, y coloreamos la celda por una regla simple:
        # - si hay 1 sola materia activa ese día, se usa su color
        # - si hay varias, se usa el primer color de la paleta
        day_color: Optional[str] = None

        active_subjects_today = 0

        for subj in subjects:
            base = int(per_subject_base.get(subj, 0))
            if base <= 0:
                continue

            qty = _round_half_up(base * mult)
            if qty <= 0:
                continue

            cat = subj_catalog[subj]
            p = prog_by_subject[subj]

            # avanzar cuadernillos agotados
            while p.booklet_idx < len(cat):
                total_q = int(cat[p.booklet_idx]["total_q"])
                if p.pos_in_booklet >= total_q:
                    p.booklet_idx += 1
                    p.pos_in_booklet = 0
                else:
                    break

            if p.booklet_idx >= len(cat):
                continue  # materia sin más preguntas

            booklet_name = cat[p.booklet_idx]["name"]
            total_q = int(cat[p.booklet_idx]["total_q"])

            # asignar rango
            start_q = p.pos_in_booklet + 1
            end_q = min(p.pos_in_booklet + qty, total_q)
            real_qty = end_q - start_q + 1

            blocks.append(DayBlock(
                subject=subj,
                booklet_name=booklet_name,
                start_q=start_q,
                end_q=end_q,
                qty=real_qty,
            ))

            p.pos_in_booklet += real_qty

            # nota de completado por materia
            if p.pos_in_booklet >= total_q:
                notes.append(f"🎉 {subj}: Cuadernillo {booklet_name} COMPLETADO")

            active_subjects_today += 1

        if active_subjects_today == 1 and blocks:
            day_color = subject_color_map.get(blocks[0].subject)
        elif active_subjects_today > 1:
            day_color = palette[0]

        is_rest = (len(blocks) == 0)
        if is_rest:
            notes.append("Descanso")
            day_color = None

        days.append(PlanDay(day=d, weekday=wd, is_rest=is_rest, color=day_color, blocks=blocks, notes=notes))
        d += timedelta(days=1)

    return PlanModel(meta=cfg["meta"], days=days)


# -------- Tipo 3: monomateria por día (sin multiplicadores) --------

def _build_type3(cfg: Dict[str, Any]) -> PlanModel:
    start = date.fromisoformat(cfg["date_range"]["start"])
    end = date.fromisoformat(cfg["date_range"]["end"]) 
    active_weekdays = set(cfg["calendar_rules"]["active_weekdays"])

    subj_catalog: Dict[str, List[Dict[str, Any]]] = cfg["content"]["subjects"]

    assign = cfg["assignments"]["subject_by_weekday"]

    # materias usadas (excluye descanso)
    subjects_used = []
    for k, v in assign.items():
        if isinstance(v, str) and v.strip().lower() != "descanso":
            subjects_used.append(v)
    # unique preservando orden
    seen = set()
    subjects_used = [s for s in subjects_used if not (s in seen or seen.add(s))]

    palette = cfg["style"]["palette"]
    subject_color_map = cfg["style"].get("subject_color_map") or cycle_map(subjects_used, palette)

    prog_by_subject: Dict[str, SubjectProgress] = {s: SubjectProgress() for s in subjects_used}

    days: List[PlanDay] = []
    d = start

    while d <= end:
        wd = d.weekday()

        if wd not in active_weekdays:
            days.append(PlanDay(day=d, weekday=wd, is_rest=True, color=None, blocks=[], notes=["Descanso"],))
            d += timedelta(days=1)
            continue

        assigned = assign.get(str(wd), "Descanso")
        if not isinstance(assigned, str):
            assigned = "Descanso"

        if assigned.strip().lower() == "descanso":
            days.append(PlanDay(day=d, weekday=wd, is_rest=True, color=None, blocks=[], notes=["Descanso"],))
            d += timedelta(days=1)
            continue

        quota = daily_quota_type1_or_3(cfg, wd)
        if quota <= 0:
            days.append(PlanDay(day=d, weekday=wd, is_rest=True, color=None, blocks=[], notes=["Descanso"],))
            d += timedelta(days=1)
            continue

        subj = assigned
        if subj not in subj_catalog:
            # si asigna materia inexistente, lo tratamos como descanso para no explotar
            days.append(PlanDay(day=d, weekday=wd, is_rest=True, color=None, blocks=[], notes=[f"Descanso (materia inválida: {subj})"],))
            d += timedelta(days=1)
            continue

        p = prog_by_subject.setdefault(subj, SubjectProgress())
        cat = subj_catalog[subj]

        # avanzar cuadernillos agotados
        while p.booklet_idx < len(cat):
            total_q = int(cat[p.booklet_idx]["total_q"])
            if p.pos_in_booklet >= total_q:
                p.booklet_idx += 1
                p.pos_in_booklet = 0
            else:
                break

        if p.booklet_idx >= len(cat):
            days.append(PlanDay(day=d, weekday=wd, is_rest=True, color=None, blocks=[], notes=[f"{subj}: sin más preguntas"],))
            d += timedelta(days=1)
            continue

        booklet_name = cat[p.booklet_idx]["name"]
        total_q = int(cat[p.booklet_idx]["total_q"])

        start_q = p.pos_in_booklet + 1
        end_q = min(p.pos_in_booklet + quota, total_q)
        real_qty = end_q - start_q + 1

        blocks = [DayBlock(subject=subj, booklet_name=booklet_name, start_q=start_q, end_q=end_q, qty=real_qty)]
        notes: List[str] = []

        p.pos_in_booklet += real_qty
        if p.pos_in_booklet >= total_q:
            notes.append(f"🎉 {subj}: Cuadernillo {booklet_name} COMPLETADO")

        day_color = subject_color_map.get(subj)
        days.append(PlanDay(day=d, weekday=wd, is_rest=False, color=day_color, blocks=blocks, notes=notes))
        d += timedelta(days=1)

    return PlanModel(meta=cfg["meta"], days=days)


# ------------------------------------------------
# 7) Render Word (manteniendo tu diseño base)
# ------------------------------------------------


def _clear_cell(cell) -> None:
    # Quita párrafos y deja uno limpio
    cell.text = ""


def _add_paragraph(cell, text: str, size: int = 12, bold: bool = True, align_center: bool = True):
    p = cell.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def render_docx(cfg: Dict[str, Any], model: PlanModel, out_path: str) -> None:
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # ===== Header con logos a ambos lados =====
    branding = cfg["style"]["branding"]
    header = doc.sections[0].header

    header_table = header.add_table(rows=1, cols=3, width=Inches(6))
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    logo_path = branding.get("logo_path")

    # Izq: Logo
    cell_left = header_table.cell(0, 0)
    p_left = cell_left.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if logo_path:
        try:
            p_left.add_run().add_picture(logo_path, width=Inches(0.8))
        except Exception:
            p_left.add_run(" ")

    # Centro: título + subtítulo
    cell_center = header_table.cell(0, 1)
    p_center = cell_center.paragraphs[0]
    run_center = p_center.add_run(branding.get("header_title", "📚 CALENDARIO ICFES 📚"))
    run_center.bold = True
    run_center.font.size = Pt(14)
    p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = cell_center.add_paragraph(branding.get("header_subtitle", "- Preicfes Material -"))
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if subtitle.runs:
        subtitle.runs[0].font.size = Pt(10)

    # Der: Logo
    cell_right = header_table.cell(0, 2)
    p_right = cell_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if logo_path:
        try:
            p_right.add_run().add_picture(logo_path, width=Inches(0.8))
        except Exception:
            p_right.add_run(" ")

    # ===== Portada =====
    student = cfg["meta"]["student_name"]
    plan_label = cfg["meta"].get("plan_label", "ICFES")

    title_text = f"📚✨ Plan {plan_label} – {student} ✨📚"
    p_title = doc.add_paragraph()
    r_title = p_title.add_run(title_text)
    r_title.bold = True
    r_title.font.name = "Arial Black"
    r_title.font.size = Pt(20)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    start = date.fromisoformat(cfg["date_range"]["start"])
    end = date.fromisoformat(cfg["date_range"]["end"])

    intro = doc.add_paragraph(
        f"🎯 Este Plan de Estudios ha sido diseñado especialmente para {student}."
        f"Comienza el {start.strftime('%d/%m/%Y')} y va hasta el {end.strftime('%d/%m/%Y')}. 📅"
        "📌 Objetivo: resolver preguntas de forma estratégica para mejorar tu rendimiento en el ICFES. 💪"
        "✅ Recomendación: 2 minutos por pregunta + 1 minuto para analizar la explicación (3 min total)."
        "Si no entiendes un tema, apóyate en temarios/clases del Drive o pídeme explicación aquí. 🤖📚"
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Bloque de cuadernillos (Tipo 1) o materias (Tipo 2/3)
    pt = int(cfg["meta"]["plan_type"])

    if pt == 1:
        names = [b["name"] for b in cfg["content"]["global_booklets"]]
        emoji_map = cfg["style"].get("booklet_emoji_map", {})
        lines = []
        for n in names:
            em = emoji_map.get(n, "📘")
            lines.append(f"{em} {n}")
        doc.add_paragraph("".join(lines))
    else:
        subjects = list(cfg["content"]["subjects"].keys())
        doc.add_paragraph("📚 Materias del plan:" + "".join([f"🔹 {s}" for s in subjects]))

    # Personalización (extras)
    extras = cfg.get("extras", {})
    p_sub = doc.add_paragraph()
    r_sub = p_sub.add_run("🔧 Personalización")
    r_sub.bold = True
    r_sub.font.size = Pt(14)
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    focus = extras.get("focus_subjects", [])
    weekly_extra_q = extras.get("weekly_extra_q", 0)

    p_focus = doc.add_paragraph()
    r_focus = p_focus.add_run("🧠 Materias con mayor intensidad:")
    r_focus.bold = True
    r_focus.font.size = Pt(14)

    if focus:
        for m in focus:
            doc.add_paragraph(f"🔹 {m}", style="List Bullet")
        if weekly_extra_q:
            doc.add_paragraph(
                f"🚀 Cada semana resolverás como extra para reforzar {weekly_extra_q} preguntas de tu elección de los ✨ Compilados VIP."
            )
    else:
        doc.add_paragraph("Preferiste darle intensidad normal a todas 😊")

    # Recomendaciones
    doc.add_heading('📌 Recomendaciones Importantes', level=2)
    recs = [
        "📩 Completa y envía tus respuestas en los formularios",
        "🔍 Revisa y analiza la explicación de cada respuesta",
        "📊 Registra tus resultados en tu Registro de Progreso ICFES",
        "🖨️ Imprime tu Plan de Estudios y tu Registro",
    ]
    for rec in recs:
        doc.add_paragraph(rec, style='List Bullet')

    doc.add_page_break()

    # ===== Calendario (semanas, tabla 2×7) =====
    header_fill = cfg["style"]["ui"].get("header_fill", "1F4E79")
    checkbox = cfg["style"]["ui"].get("checkbox_unchecked", "⬜")

    # Index por fecha para acceso rápido
    day_by_date: Dict[date, PlanDay] = {pd.day: pd for pd in model.days}

    week_num = 1
    page_break_every_weeks = int(cfg["style"]["ui"].get("page_break_every_weeks", 2))

    for week in iter_weeks_monday_to_sunday(start, end):
        # salto de página controlado
        if week_num > 1 and page_break_every_weeks > 0 and (week_num - 1) % page_break_every_weeks == 0:
            doc.add_page_break()

        monday = week[0]
        doc.add_heading(f"🗓️ Semana #{week_num} - Del {monday.strftime('%d/%m/%Y')}", level=2)

        table = doc.add_table(rows=2, cols=7)
        table.style = 'Table Grid'
        table.autofit = True
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # encabezados
        for i in range(7):
            cell = table.rows[0].cells[i]
            run = cell.paragraphs[0].add_run(DAYS_ES[i].capitalize())
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255, 255, 255)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            shade_cell(cell, header_fill)

        # celdas del contenido
        for i in range(7):
            cell = table.rows[1].cells[i]
            current = week[i]

            if current < start or current > end:
                cell.text = ""
                continue

            pd = day_by_date.get(current)
            _clear_cell(cell)

            if pd is None:
                _add_paragraph(cell, f"📅 {current.strftime('%d/%m/%Y')}")
                _add_paragraph(cell, "Descanso", size=12, bold=True)
                continue

            # Fecha
            _add_paragraph(cell, f"📅 {current.strftime('%d/%m/%Y')}")
            _add_paragraph(cell, "📝 Tareas:")

            if pd.is_rest or not pd.blocks:
                _add_paragraph(cell, "Descanso", size=12, bold=True)
            else:
                # ✅ NUEVO: Separadores tipo rectángulo por materia
                # Solo para Tipo 1 y Tipo 2 (en Tipo 3 siempre es una sola materia)
                if pt in (1, 2) and len(pd.blocks) > 0:
                    inner = cell.add_table(rows=len(pd.blocks), cols=1)
                    inner.style = 'Table Grid'
                    inner.autofit = True
                    for r_i, b in enumerate(pd.blocks):
                        c = inner.cell(r_i, 0)
                        c.text = ""
                        p = c.paragraphs[0]
                        rr = p.add_run(f"{b.subject}: {b.qty} ({b.start_q}–{b.end_q})")
                        rr.bold = True
                        rr.font.size = Pt(11)
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    # Tipo 3 (o fallback): una sola línea
                    for b in pd.blocks:
                        _add_paragraph(cell, f"- {b.qty} de {b.subject} ({b.start_q}–{b.end_q})", size=12, bold=True)

            # notas
            if pd.notes:
                cell.add_paragraph("")
                for n in pd.notes:
                    _add_paragraph(cell, n, size=11, bold=True)

            cell.add_paragraph("")
            _add_paragraph(cell, checkbox, size=12, bold=True)

            # color de celda
            if pd.color:
                try:
                    shade_cell(cell, pd.color.lstrip("#"))
                except Exception:
                    pass

        week_num += 1

    doc.save(out_path)


# ------------------------------------------------
# 8) Orquestador (lo que llaman los runners)
# ------------------------------------------------

def generate_plan_docx(cfg: Dict[str, Any], out_filename: str) -> str:
    model = build_plan_model(cfg)
    render_docx(cfg, model, out_filename)
    return out_filename


def colab_download(path: str) -> None:
    if files is None:
        print(f"Archivo generado: {path}")
        print("(No estás en Colab, descarga manualmente.)")
        return
    files.download(path)


# ==============================
# ===== END CORE (no correr) ===
# ==============================

